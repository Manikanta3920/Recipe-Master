import streamlit as st
import os, time, random, io
from dotenv import load_dotenv
from google import genai
from google.genai.errors import ClientError
from fpdf import FPDF
from docx import Document

# ---------------- CONFIG ----------------
load_dotenv()
client = genai.Client(api_key=os.getenv("GOOGLE_API_KEY"))
MODEL = "models/gemini-2.5-flash"

st.set_page_config(
    page_title="RecipeMaster 🍽️",
    page_icon="🍲",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ---------------- HELPERS ----------------
def clean_text_for_pdf(text: str) -> str:
    replacements = {
        "–": "-", "—": "-", "“": '"', "”": '"',
        "‘": "'", "’": "'", "₹": "Rs.", "•": "-"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text

def difficulty_color(level):
    return {
        "Easy": "#10b981",  # Emerald
        "Medium": "#f59e0b",  # Amber
        "Hard": "#ef4444"  # Red
    }[level]

def difficulty_icon(level):
    return {
        "Easy": "🌱",
        "Medium": "🔥",
        "Hard": "👨‍🍳"
    }[level]

# ---------------- THEME & STYLING ----------------
dark_mode = st.toggle("🌙 Dark Mode", value=False)

# Color Palette
bg = "#0f172a" if dark_mode else "#f8fafc"
text = "#f1f5f9" if dark_mode else "#0f172a"
card = "#1e293b" if dark_mode else "#ffffff"
accent = "#f97316"  # Orange-500
primary = "#3b82f6"  # Blue-500
surface = "#334155" if dark_mode else "#f1f5f9"

# Custom CSS with enhanced styling
st.markdown(f"""
<style>
/* Base Styles */
.stApp {{
    background-color: {bg};
    color: {text};
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
}}

/* Custom Card */
.custom-card {{
    background-color: {card};
    padding: 2rem;
    border-radius: 1.5rem;
    border: 1px solid {surface};
    box-shadow: 0 10px 40px rgba(0,0,0,0.1);
    margin-bottom: 1.5rem;
    transition: transform 0.2s ease, box-shadow 0.2s ease;
}}

.custom-card:hover {{
    transform: translateY(-2px);
    box-shadow: 0 20px 60px rgba(0,0,0,0.15);
}}

/* Badge Styles */
.difficulty-badge {{
    display: inline-flex;
    align-items: center;
    padding: 0.5rem 1rem;
    border-radius: 9999px;
    font-weight: 600;
    font-size: 0.875rem;
    margin: 0.5rem 0;
    gap: 0.5rem;
}}

/* Button Styles */
.stButton > button {{
    background: linear-gradient(135deg, {accent}, #fb923c);
    color: white;
    font-weight: 600;
    font-size: 1rem;
    border: none;
    border-radius: 1rem;
    padding: 1rem 2rem;
    transition: all 0.3s ease;
    width: 100%;
    margin-top: 1rem;
}}

.stButton > button:hover {{
    transform: translateY(-2px);
    box-shadow: 0 10px 25px rgba(249, 115, 22, 0.3);
    background: linear-gradient(135deg, #ea580c, {accent});
}}

/* Input Styles */
.stTextInput > div > div > input {{
    border: 2px solid {surface};
    border-radius: 1rem;
    padding: 0.75rem 1rem;
    font-size: 1rem;
    transition: all 0.3s ease;
    background-color: {card};
    color: {text};
}}

.stTextInput > div > div > input:focus {{
    border-color: {accent};
    box-shadow: 0 0 0 3px rgba(249, 115, 22, 0.1);
}}

/* Selectbox Styles */
.stSelectbox > div > div {{
    border: 2px solid {surface};
    border-radius: 1rem;
    background-color: {card};
}}

/* Slider Styles */
.stSlider > div > div > div {{
    background: linear-gradient(90deg, {accent}, #fb923c);
}}

/* Progress Bar */
.stProgress > div > div > div {{
    background: linear-gradient(90deg, {primary}, {accent});
}}

/* Expander Styles */
.streamlit-expanderHeader {{
    background-color: {surface} !important;
    border-radius: 1rem !important;
    border: 1px solid {surface} !important;
}}

.streamlit-expanderContent {{
    background-color: {card};
    border-radius: 0 0 1rem 1rem;
    padding: 1.5rem;
}}

/* Download Buttons */
.download-btn {{
    background: linear-gradient(135deg, {primary}, #60a5fa) !important;
    margin: 0.5rem 0;
}}

/* Status Messages */
.stAlert {{
    border-radius: 1rem;
    border: none;
}}

/* Custom Divider */
.custom-divider {{
    height: 3px;
    background: linear-gradient(90deg, transparent, {accent}, transparent);
    margin: 2rem 0;
    border: none;
}}

/* Live Preview Card */
.preview-card {{
    background: linear-gradient(135deg, {surface}, {card});
    border-left: 4px solid {accent};
    padding: 1.5rem;
    border-radius: 1rem;
    margin: 1rem 0;
}}

/* Typography */
h1, h2, h3 {{
    font-weight: 700 !important;
    letter-spacing: -0.025em;
}}

.caption {{
    color: {text} !important;
    opacity: 0.7;
    font-size: 0.9rem;
}}

/* Animation for recipe reveal */
@keyframes fadeIn {{
    from {{ opacity: 0; transform: translateY(10px); }}
    to {{ opacity: 1; transform: translateY(0); }}
}}

.recipe-line {{
    animation: fadeIn 0.3s ease forwards;
    opacity: 0;
    margin: 0.5rem 0;
}}

/* Loading Animation */
@keyframes pulse {{
    0%, 100% {{ opacity: 1; }}
    50% {{ opacity: 0.5; }}
}}

.loading-pulse {{
    animation: pulse 1.5s ease-in-out infinite;
}}
</style>
""", unsafe_allow_html=True)

# ---------------- HEADER ----------------
st.markdown("""
<div style="text-align: center; padding: 2rem 0;">
    <h1 style="font-size: 3.5rem; margin-bottom: 0.5rem;">
        🍽️ RecipeMaster
    </h1>
    <div class="caption" style="font-size: 1.1rem; margin-bottom: 1rem;">
        Transform ingredients into culinary masterpieces with AI
    </div>
</div>
""", unsafe_allow_html=True)

# ---------------- HERO SECTION ----------------
col1, col2, col3 = st.columns([1, 2, 1])

with col2:
    st.markdown("""
    <div class="custom-card" style="text-align: center;">
        <div style="font-size: 4rem; margin-bottom: 1rem;">✨</div>
        <h3 style="margin-bottom: 1rem;">AI-Powered Recipe Generation</h3>
        <p class="caption">
            Get detailed recipes with exact measurements, step-by-step instructions, 
            and professional tips in seconds. Perfect for home cooks and chefs alike.
        </p>
    </div>
    """, unsafe_allow_html=True)

# ---------------- INPUT SECTION ----------------
st.markdown("<div class='custom-divider'></div>", unsafe_allow_html=True)
st.markdown("<h2 style='margin-bottom: 1.5rem;'>🎯 Recipe Parameters</h2>", unsafe_allow_html=True)

# Input Card
st.markdown("<div class='custom-card'>", unsafe_allow_html=True)

col1, col2 = st.columns(2)

with col1:
    topic = st.text_input(
        "**🍽️ Recipe Name**",
        placeholder="Enter dish name (e.g., Paneer Butter Masala, Chocolate Cake)",
        help="Be specific for best results!"
    )

with col2:
    difficulty = st.selectbox(
        "**🧑‍🍳 Difficulty Level**",
        ["Easy", "Medium", "Hard"],
        index=0,
        help="Choose based on your cooking experience"
    )

# Word Count Slider with better styling
st.markdown("<div style='margin: 2rem 0;'>", unsafe_allow_html=True)
word_count = st.slider(
    "**📝 Recipe Detail Level**",
    min_value=200,
    max_value=1000,
    value=500,
    step=50,
    help="Longer recipes include more tips, variations, and detailed instructions"
)

# Visual indicator for word count
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    st.markdown(f"""
    <div style="text-align: center; margin: 1rem 0; padding: 1rem; 
                background: {surface}; border-radius: 1rem;">
        <div class="caption">Target Length</div>
        <div style="font-size: 2rem; font-weight: 700; color: {accent};">
            {word_count}
        </div>
        <div class="caption">words</div>
    </div>
    """, unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)

# Difficulty Badge
st.markdown(f"""
<div class="difficulty-badge" style="background-color: {difficulty_color(difficulty)}20; 
                                      border: 2px solid {difficulty_color(difficulty)};">
    <span style="font-size: 1.2rem;">{difficulty_icon(difficulty)}</span>
    <span style="color: {difficulty_color(difficulty)};">{difficulty} Level</span>
</div>
""", unsafe_allow_html=True)

st.markdown("</div>", unsafe_allow_html=True)

# ---------------- GENERATE BUTTON ----------------
generate_btn = st.button(
    "🚀 **Generate Recipe**",
    type="primary",
    use_container_width=True
)

# ---------------- LIVE PREVIEW ----------------
if topic:
    st.markdown("<div class='preview-card'>", unsafe_allow_html=True)
    st.markdown(f"""
    <div style="display: flex; align-items: center; gap: 1rem;">
        <div style="font-size: 1.5rem;">🔍</div>
        <div>
            <div style="font-weight: 600; font-size: 1.1rem;">Live Preview</div>
            <div class="caption">
                Generating <span style="color: {accent}; font-weight: 600;">"{topic}"</span> 
                as a <span style="color: {difficulty_color(difficulty)}; font-weight: 600;">{difficulty}</span> 
                level recipe
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

# ---------------- GENERATION SECTION ----------------
if generate_btn and topic:
    # Loading Animation
    with st.container():
        st.markdown("<div class='custom-card'>", unsafe_allow_html=True)
        
        # Animated progress bar
        progress_bar = st.progress(0, text="Starting recipe generation...")
        
        # Loading messages with icons
        loading_stages = [
            {"icon": "🔍", "text": "Researching culinary techniques..."},
            {"icon": "📚", "text": "Gathering ingredient insights..."},
            {"icon": "🧠", "text": "Consulting AI chef..."},
            {"icon": "🍳", "text": "Crafting cooking instructions..."},
            {"icon": "✨", "text": "Adding final touches..."}
        ]
        
        status_container = st.empty()
        
        # Simulated loading with stages
        for i in range(5):
            stage = loading_stages[i]
            progress = (i + 1) * 20
            progress_bar.progress(progress, 
                                 text=f"{stage['icon']} {stage['text']}")
            status_container.markdown(f"""
            <div class="loading-pulse" style="text-align: center; padding: 1rem;">
                <div style="font-size: 2rem;">{stage['icon']}</div>
                <div class="caption">{stage['text']}</div>
            </div>
            """, unsafe_allow_html=True)
            time.sleep(0.5)
        
        # Generate recipe
        prompt = f"""
        Create a {difficulty.lower()} level food recipe blog for "{topic}".

        Include:
        - Catchy title
        - Short introduction
        - Serves, prep time, cook time
        - Ingredients list
        - Step-by-step cooking instructions
        - Tips & variations
        - Conclusion

        Tone: friendly, professional.
        Length: about {word_count} words.
        """
        
        try:
            response = client.models.generate_content(
                model=MODEL,
                contents=prompt
            )
            
            recipe_text = response.text
            st.session_state["recipe"] = recipe_text
            
            # Clear loading indicators
            progress_bar.empty()
            status_container.empty()
            
            st.markdown("</div>", unsafe_allow_html=True)
            
            # Success Message
            st.markdown(f"""
            <div style="text-align: center; padding: 2rem; 
                        background: linear-gradient(135deg, {difficulty_color(difficulty)}20, transparent);
                        border-radius: 1.5rem; margin: 2rem 0;">
                <div style="font-size: 3rem;">🎉</div>
                <h2>Recipe Generated Successfully!</h2>
                <p class="caption">Your AI-crafted recipe is ready to explore</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Recipe Display
            with st.expander("📖 **View Full Recipe**", expanded=True):
                st.markdown("<div style='padding: 1.5rem; background-color: rgba(0,0,0,0.02); border-radius: 1rem;'>", 
                           unsafe_allow_html=True)
                
                # Animated line-by-line reveal
                lines = recipe_text.split("\n")
                container = st.container()
                with container:
                    for i, line in enumerate(lines):
                        if line.strip():
                            st.markdown(
                                f'<div class="recipe-line" style="animation-delay: {i*0.03}s;">{line}</div>',
                                unsafe_allow_html=True
                            )
                            time.sleep(0.01)
                
                st.markdown("</div>", unsafe_allow_html=True)
            
            # ---------------- DOWNLOAD SECTION ----------------
            st.markdown("<div class='custom-divider'></div>", unsafe_allow_html=True)
            st.markdown("<h2 style='margin-bottom: 1.5rem;'>📥 Export Recipe</h2>", unsafe_allow_html=True)
            
            st.markdown("""
            <div class="custom-card">
                <div style="text-align: center; margin-bottom: 1.5rem;">
                    <div style="font-size: 2rem;">💾</div>
                    <h3>Save Your Recipe</h3>
                    <p class="caption">Download in your preferred format for easy access</p>
                </div>
            """, unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns(3)
            
            # PDF Generation
            safe_text = clean_text_for_pdf(recipe_text)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=11)
            
            for line in safe_text.split("\n"):
                pdf.multi_cell(0, 8, line)
            
            pdf_bytes = pdf.output(dest="S").encode("latin-1")
            
            with col1:
                st.download_button(
                    label="**📄 Download PDF**",
                    data=pdf_bytes,
                    file_name=f"RecipeMaster_{topic.replace(' ', '_')}.pdf",
                    mime="application/pdf",
                    help="Best for printing and sharing",
                    use_container_width=True
                )
            
            # DOCX Generation
            doc = Document()
            doc.add_heading(topic, level=1)
            doc.add_paragraph(recipe_text)
            buffer = io.BytesIO()
            doc.save(buffer)
            
            with col2:
                st.download_button(
                    label="**📝 Download DOCX**",
                    data=buffer.getvalue(),
                    file_name=f"RecipeMaster_{topic.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Editable in Microsoft Word",
                    use_container_width=True
                )
            
            # Text File
            with col3:
                st.download_button(
                    label="**📱 Download TXT**",
                    data=recipe_text,
                    file_name=f"RecipeMaster_{topic.replace(' ', '_')}.txt",
                    mime="text/plain",
                    help="Simple text file for quick reference",
                    use_container_width=True
                )
            
            st.markdown("</div>", unsafe_allow_html=True)
            
        except ClientError as e:
            progress_bar.empty()
            status_container.empty()
            st.markdown("</div>", unsafe_allow_html=True)
            
            if "RESOURCE_EXHAUSTED" in str(e):
                st.error("""
                ### 🚫 Rate Limit Reached
                
                You've reached the free-tier quota. Please:
                - Wait 60 seconds before trying again
                - Try a shorter recipe
                - Consider reducing the word count
                """)
            else:
                st.error("""
                ### ⚠️ Generation Error
                
                An unexpected error occurred. Please:
                - Check your internet connection
                - Try again in a moment
                - Contact support if issue persists
                """)

elif generate_btn:
    st.warning("""
    ### ⚠️ Missing Recipe Name
    
    Please enter a recipe name to get started!
    """)

# ---------------- FOOTER ----------------
st.markdown("<div class='custom-divider'></div>", unsafe_allow_html=True)

st.markdown(f"""
<div style="text-align: center; padding: 2rem 0; color: {text}; opacity: 0.7;">
    <div style="font-size: 1.5rem; margin-bottom: 0.5rem;">
        🍳 ✨ 👨‍🍳
    </div>
    <div style="margin-bottom: 0.5rem;">
        <strong>RecipeMaster</strong> • Powered by Google Gemini Flash
    </div>
    <div class="caption">
        Crafted with ❤️ for food enthusiasts everywhere
    </div>
</div>
""", unsafe_allow_html=True)
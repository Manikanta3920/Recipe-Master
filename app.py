import streamlit as st
import os, time, random, io

from dotenv import load_dotenv
from google import genai
from google.genai.errors import ClientError

from fpdf import FPDF
from docx import Document


# =====================================================
# ✅ CONFIG (WORKS FOR LOCAL + RENDER DEPLOYMENT)
# =====================================================

# Load .env only for local development
load_dotenv()

# Render will use Environment Variables (not .env file)
api_key = os.getenv("GEMINI_API_KEY")

# If key not found → stop app
if not api_key:
    st.error("❌ GOOGLE_API_KEY not found!")
    st.info("➡️ Please add GOOGLE_API_KEY in Render → Environment Variables.")
    st.stop()

# ✅ Create Gemini Client properly
client = genai.Client(api_key=api_key)

# Gemini Model
MODEL = "models/gemini-2.5-flash"


# =====================================================
# ✅ STREAMLIT PAGE SETTINGS
# =====================================================

st.set_page_config(
    page_title="RecipeMaster 🍽️",
    page_icon="🍲",
    layout="wide",
    initial_sidebar_state="collapsed"
)


# =====================================================
# ✅ HELPERS
# =====================================================

def clean_text_for_pdf(text: str) -> str:
    replacements = {
        "–": "-", "—": "-", "“": '"', "”": '"',
        "‘": "'", "’": "'", "₹": "Rs.", "•": "-"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def difficulty_color(level):
    return {
        "Easy": "#10b981",
        "Medium": "#f59e0b",
        "Hard": "#ef4444"
    }[level]


def difficulty_icon(level):
    return {
        "Easy": "🌱",
        "Medium": "🔥",
        "Hard": "👨‍🍳"
    }[level]


# =====================================================
# ✅ THEME & STYLING (Your Same Code)
# =====================================================

dark_mode = st.toggle("🌙 Dark Mode", value=False)

bg = "#0f172a" if dark_mode else "#f8fafc"
text = "#f1f5f9" if dark_mode else "#0f172a"
card = "#1e293b" if dark_mode else "#ffffff"
accent = "#f97316"
primary = "#3b82f6"
surface = "#334155" if dark_mode else "#f1f5f9"

st.markdown(f"""
<style>
.stApp {{
    background-color: {bg};
    color: {text};
}}
</style>
""", unsafe_allow_html=True)


# =====================================================
# ✅ HEADER
# =====================================================

st.markdown("""
<div style="text-align: center; padding: 2rem 0;">
    <h1 style="font-size: 3.5rem;">🍽️ RecipeMaster</h1>
    <p>Transform ingredients into culinary masterpieces with AI</p>
</div>
""", unsafe_allow_html=True)


# =====================================================
# ✅ INPUT SECTION
# =====================================================

topic = st.text_input("🍽️ Recipe Name", placeholder="Paneer Butter Masala")
difficulty = st.selectbox("🧑‍🍳 Difficulty", ["Easy", "Medium", "Hard"])
word_count = st.slider("📝 Detail Level", 200, 1000, 500, step=50)

generate_btn = st.button("🚀 Generate Recipe")


# =====================================================
# ✅ GENERATION SECTION
# =====================================================

if generate_btn and topic:

    prompt = f"""
    Create a {difficulty.lower()} level food recipe blog for "{topic}".

    Include:
    - Catchy title
    - Short introduction
    - Serves, prep time, cook time
    - Ingredients list
    - Step-by-step cooking instructions
    - Tips & variations
    - Conclusion

    Tone: friendly, professional.
    Length: about {word_count} words.
    """

    try:
        # ✅ Correct Gemini Call
        response = client.models.generate_content(
            model=MODEL,
            contents=prompt
        )

        recipe_text = response.text

        st.success("🎉 Recipe Generated Successfully!")
        st.markdown(recipe_text)

        # ---------------- DOWNLOAD SECTION ----------------

        st.subheader("📥 Download Recipe")

        # PDF Download
        safe_text = clean_text_for_pdf(recipe_text)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)

        for line in safe_text.split("\n"):
            pdf.multi_cell(0, 8, line)

        pdf_bytes = pdf.output(dest="S").encode("latin-1")

        st.download_button(
            label="📄 Download PDF",
            data=pdf_bytes,
            file_name=f"{topic}.pdf",
            mime="application/pdf"
        )

        # DOCX Download
        doc = Document()
        doc.add_heading(topic, level=1)
        doc.add_paragraph(recipe_text)

        buffer = io.BytesIO()
        doc.save(buffer)

        st.download_button(
            label="📝 Download DOCX",
            data=buffer.getvalue(),
            file_name=f"{topic}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # TXT Download
        st.download_button(
            label="📱 Download TXT",
            data=recipe_text,
            file_name=f"{topic}.txt",
            mime="text/plain"
        )

    except ClientError as e:
        st.error("⚠️ Gemini API Error!")
        st.write(e)

elif generate_btn:
    st.warning("⚠️ Please enter a recipe name first!")
